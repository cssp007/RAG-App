# services/document_service.py

import hashlib
import os

from pypdf import PdfReader
from docx import Document


class DocumentService:


    # ==================================
    # INITIALIZE DOCUMENT SERVICE
    # ==================================

    def __init__(
        self,
        vector_service,
        embedding_service
    ):

        self.vector_service = (
            vector_service
        )

        self.embedding_service = (
            embedding_service
        )


    # ==================================
    # PROCESS DOCUMENT
    # ==================================

    def process_document(
        self,
        file_path,
        original_filename,
        document_id
    ):

        # ------------------------------
        # CHECK FILE EXISTS
        # ------------------------------

        if not os.path.exists(
            file_path
        ):

            raise FileNotFoundError(

                "Uploaded file could not be found."

            )


        # ------------------------------
        # CREATE FILE HASH
        # ------------------------------

        file_hash = (

            self.calculate_file_hash(
                file_path
            )

        )


        # ------------------------------
        # CHECK DUPLICATE FILE
        # ------------------------------

        existing_document = (

            self.vector_service
            .get_document_by_hash(
                file_hash
            )

        )


        if existing_document:

            existing_filename = (

                existing_document.get(
                    "filename",
                    original_filename
                )

            )


            existing_document_id = (

                existing_document.get(
                    "document_id",
                    ""
                )

            )


            return {

                "status":
                "duplicate",

                "message":

                f"Document '{original_filename}' "
                f"has already been uploaded.",

                "document_id":
                existing_document_id,

                "filename":
                existing_filename

            }


        # ------------------------------
        # EXTRACT TEXT
        # ------------------------------

        text = (

            self.extract_text(
                file_path
            )

        )


        if not text:

            raise ValueError(

                "No readable text was found "
                "in the uploaded document."

            )


        # ------------------------------
        # CREATE TEXT CHUNKS
        # ------------------------------

        chunks = (

            self.chunk_text(
                text
            )

        )


        if not chunks:

            raise ValueError(

                "Unable to create text chunks "
                "from the uploaded document."

            )


        # ------------------------------
        # CREATE EMBEDDINGS
        # ------------------------------

        embeddings = (

            self.embedding_service
            .create_embeddings(
                chunks
            )

        )


        if not embeddings:

            raise ValueError(

                "Unable to create embeddings "
                "for the uploaded document."

            )


        # ------------------------------
        # CREATE CHUNK IDS
        # ------------------------------

        chunk_ids = []


        for index in range(
            len(
                chunks
            )
        ):

            chunk_id = (

                f"{document_id}_chunk_{index}"

            )

            chunk_ids.append(
                chunk_id
            )


        # ------------------------------
        # CREATE METADATA
        # ------------------------------

        metadatas = []


        for index in range(
            len(
                chunks
            )
        ):

            metadata = {

                "document_id":
                document_id,

                "filename":
                original_filename,

                "file_hash":
                file_hash,

                "chunk_id":
                index,

                "total_chunks":
                len(
                    chunks
                )

            }


            metadatas.append(
                metadata
            )


        # ------------------------------
        # STORE IN VECTOR DATABASE
        # ------------------------------

        self.vector_service.add_document_chunks(

            ids=
            chunk_ids,

            documents=
            chunks,

            embeddings=
            embeddings,

            metadatas=
            metadatas

        )


        return {

            "status":
            "success",

            "message":

            f"Document '{original_filename}' "
            f"uploaded successfully.",

            "document_id":
            document_id,

            "filename":
            original_filename,

            "chunks":
            len(
                chunks
            )

        }


    # ==================================
    # CALCULATE FILE HASH
    # ==================================

    def calculate_file_hash(
        self,
        file_path
    ):

        sha256 = (
            hashlib.sha256()
        )


        with open(
            file_path,
            "rb"
        ) as file:

            while True:

                data = (
                    file.read(
                        8192
                    )
                )


                if not data:

                    break


                sha256.update(
                    data
                )


        return (
            sha256.hexdigest()
        )


    # ==================================
    # EXTRACT TEXT
    # ==================================

    def extract_text(
        self,
        file_path
    ):

        extension = (

            os.path.splitext(
                file_path
            )[1]

            .lower()

        )


        # ------------------------------
        # TXT FILE
        # ------------------------------

        if extension == ".txt":

            return (

                self.extract_text_from_txt(
                    file_path
                )

            )


        # ------------------------------
        # PDF FILE
        # ------------------------------

        if extension == ".pdf":

            return (

                self.extract_text_from_pdf(
                    file_path
                )

            )


        # ------------------------------
        # DOCX FILE
        # ------------------------------

        if extension == ".docx":

            return (

                self.extract_text_from_docx(
                    file_path
                )

            )


        # ------------------------------
        # DOC FILE
        # ------------------------------

        if extension == ".doc":

            raise ValueError(

                "DOC files are currently not "
                "supported for text extraction. "
                "Please convert the file to DOCX, "
                "PDF, or TXT."

            )


        raise ValueError(

            f"Unsupported file type: "
            f"{extension}"

        )


    # ==================================
    # EXTRACT TEXT FROM TXT
    # ==================================

    def extract_text_from_txt(
        self,
        file_path
    ):

        encodings = [

            "utf-8",

            "utf-8-sig",

            "latin-1"

        ]


        for encoding in encodings:

            try:

                with open(

                    file_path,

                    "r",

                    encoding=
                    encoding,

                    errors=
                    "ignore"

                ) as file:

                    return (
                        file.read()
                    )


            except UnicodeDecodeError:

                continue


        raise ValueError(

            "Unable to read TXT file."

        )


    # ==================================
    # EXTRACT TEXT FROM PDF
    # ==================================

    def extract_text_from_pdf(
        self,
        file_path
    ):

        reader = (

            PdfReader(
                file_path
            )

        )


        text_parts = []


        for page in reader.pages:

            try:

                page_text = (

                    page.extract_text()
                    or ""
                )


                if page_text:

                    text_parts.append(
                        page_text
                    )


            except Exception:

                continue


        return (

            "\n\n".join(
                text_parts
            )

        )


    # ==================================
    # EXTRACT TEXT FROM DOCX
    # ==================================

    def extract_text_from_docx(
        self,
        file_path
    ):

        document = (

            Document(
                file_path
            )

        )


        text_parts = []


        # ------------------------------
        # EXTRACT PARAGRAPHS
        # ------------------------------

        for paragraph in (
            document.paragraphs
        ):

            text = (

                paragraph.text
                .strip()

            )


            if text:

                text_parts.append(
                    text
                )


        # ------------------------------
        # EXTRACT TABLES
        # ------------------------------

        for table in (
            document.tables
        ):

            for row in table.rows:

                row_values = []


                for cell in row.cells:

                    cell_text = (

                        cell.text
                        .strip()

                    )


                    if cell_text:

                        row_values.append(
                            cell_text
                        )


                if row_values:

                    text_parts.append(

                        " | ".join(
                            row_values
                        )

                    )


        return (

            "\n\n".join(
                text_parts
            )

        )


    # ==================================
    # CLEAN TEXT
    # ==================================

    def clean_text(
        self,
        text
    ):

        if not text:

            return ""


        # Normalize line endings.

        text = (

            text.replace(
                "\r\n",
                "\n"
            )

        )


        text = (

            text.replace(
                "\r",
                "\n"
            )

        )


        # Remove repeated spaces.

        lines = []


        for line in (
            text.split(
                "\n"
            )
        ):

            line = (

                " ".join(
                    line.split()
                )

            )


            if line:

                lines.append(
                    line
                )


        return (

            "\n".join(
                lines
            )

        )


    # ==================================
    # CREATE TEXT CHUNKS
    # ==================================

    def chunk_text(
        self,
        text,
        chunk_size=1000,
        chunk_overlap=200
    ):

        text = (

            self.clean_text(
                text
            )

        )


        if not text:

            return []


        chunks = []

        start = 0

        text_length = (
            len(
                text
            )
        )


        while start < text_length:

            end = min(

                start +
                chunk_size,

                text_length

            )


            # Try to end the chunk
            # at a natural boundary.

            if end < text_length:

                boundary = max(

                    text.rfind(
                        "\n",
                        start,
                        end
                    ),

                    text.rfind(
                        ". ",
                        start,
                        end
                    ),

                    text.rfind(
                        "? ",
                        start,
                        end
                    ),

                    text.rfind(
                        "! ",
                        start,
                        end
                    )

                )


                if (

                    boundary
                    >
                    start
                    +
                    (
                        chunk_size
                        // 2
                    )

                ):

                    end = (
                        boundary + 1
                    )


            chunk = (

                text[
                    start:end
                ]
                .strip()

            )


            if chunk:

                chunks.append(
                    chunk
                )


            # Stop when the complete
            # document has been processed.

            if end >= text_length:

                break


            # Move forward while keeping
            # overlap between chunks.

            start = max(

                end -
                chunk_overlap,

                start + 1

            )


        return (
            chunks
        )


    # ==================================
    # GET ALL UNIQUE DOCUMENTS
    # ==================================

    def get_documents(
        self
    ):

        documents = (

            self.vector_service
            .get_unique_documents()

        )


        formatted_documents = []


        for document in documents:

            if not document:

                continue


            formatted_documents.append(

                {

                    "document_id":

                    document.get(
                        "document_id"
                    ),

                    "filename":

                    document.get(
                        "filename",
                        "Unknown Document"
                    ),

                    "total_chunks":

                    document.get(
                        "total_chunks",
                        0
                    )

                }

            )


        # Sort by filename.

        formatted_documents.sort(

            key=lambda item:

            item[
                "filename"
            ]
            .lower()

        )


        return (
            formatted_documents
        )


    # ==================================
    # DELETE DOCUMENT
    # ==================================

    def delete_document(
        self,
        document_id
    ):

        document = (

            self.vector_service
            .get_document(
                document_id
            )

        )


        if not document:

            raise ValueError(

                "Document not found."

            )


        filename = (

            document.get(
                "filename",
                "Document"
            )

        )


        self.vector_service.delete_document(
            document_id
        )


        return {

            "status":
            "success",

            "message":

            f"Document '{filename}' "
            f"deleted successfully."

        }


    # ==================================
    # RESET ALL DOCUMENTS
    # ==================================

    def reset_documents(
        self
    ):

        self.vector_service.reset_database()


        return {

            "status":
            "success",

            "message":
            "All documents have been removed."

        }